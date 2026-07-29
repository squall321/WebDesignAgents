# 문서 exporter 검증 — DocumentDoc 정규화·마크다운 변환·HTML/PDF/DOCX 산출과 재열기 대조
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

import pytest

from wdrender.exporter_doc import (
    blocks_from_markdown,
    export_docx,
    export_document,
    export_html,
    export_pdf,
    normalize_document,
)

ROOT = Path(__file__).resolve().parents[1]
REAL_FIXTURE = ROOT / "data" / "doc_check" / "document.json"


@pytest.fixture(scope="module")
def still(tmp_path_factory) -> Path:
    """씬 스틸 대역 이미지 1장(16:9)을 만든다 — 그림 삽입·캡션 경로 검증용."""
    from PIL import Image, ImageDraw

    path = tmp_path_factory.mktemp("assets") / "scene.png"
    img = Image.new("RGB", (960, 540), "#EBEEFA")
    d = ImageDraw.Draw(img)
    d.rectangle([60, 60, 900, 480], outline="#1428A0", width=6)
    d.text((90, 90), "SCENE STILL", fill="#101B3E")
    img.save(path)
    return path


def make_doc(still_path: Path) -> dict:
    """심의 산출을 그대로 옮긴 모양의 합성 DocumentDoc — 본문·삽화·표·요약·부록·출처."""
    return {
        "version": "1.0",
        "meta": {
            "title": "심의 문서 exporter 검증본",
            "subtitle": "영상·PPT 와 같은 심의에서 나온 세 번째 표현",
            "core_message": "보고서는 자기 심의를 각주로 단다.",
            "audience": "경영진",
            "date": "2026-07-29",
            "authors": ["스토리 아키텍트", "테크니컬 디렉터"],
            "meeting_id": "339e3bac",
            "source_report_id": 90,
        },
        "sections": [
            {
                "title": "오프닝",
                "blocks": [
                    {"type": "figure", "src": str(still_path),
                     "caption": "오프닝 씬 캡처", "alt": "오프닝"},
                    {"type": "paragraph",
                     "text": "보고서는 **한 번만** 작성합니다.", "refs": ["RA-001"]},
                    {"type": "callout", "title": "판정", "text": "가(可) — 엔진 계약 위반 없음"},
                ],
            },
            {
                "title": "근거",
                "blocks": [
                    {"type": "heading", "text": "구조 비교"},
                    # 조각 structured 원형(columns=[{key,label}] · rows=[{key: value}])
                    {"type": "table",
                     "caption": "Mount / Fork 비교",
                     "columns": [{"key": "aspect", "label": "항목"},
                                 {"key": "mount", "label": "게시"}],
                     "rows": [{"aspect": "사본 여부", "mount": "사본 아님"},
                              {"aspect": "편집 권한", "mount": "작성자 + 보직장"}],
                     "refs": ["RA-002"]},
                    {"type": "list", "ordered": True, "items": ["작성", "게시", "취합"]},
                    {"type": "quote", "text": "거부권은 행사하지 않는다.", "by": "접근성 전문가"},
                ],
            },
        ],
        "appendix": [{"title": "부록 A. 회의록",
                      "markdown": "## 결정\n\n- 게시는 사본이 아니다\n\n| 항목 | 값 |\n|---|---|\n| 턴 | 25 |\n"}],
        "footnotes": [
            {"id": "RA-001", "text": "ReportArchive 플랫폼 개요",
             "source": {"page": "1. 플랫폼 개요", "block_id": "h1_intro"}},
            {"id": "RA-002", "text": "Mount/Fork/Composite 비교",
             "source": {"page": "2. 보고서 워크플로", "block_id": "cmp_rel"}},
        ],
    }


# ── 정규화 ──────────────────────────────────────────────────────────────────

def test_normalize_기본_구조와_번호(still):
    n = normalize_document(make_doc(still))
    assert n["title"] == "심의 문서 exporter 검증본"
    assert [s["title"] for s in n["sections"]] == ["오프닝", "근거"]
    assert n["figures"] == 1
    assert n["tables"] == 2          # 본문 1 + 부록 마크다운 1
    assert n["missing_images"] == []
    # 각주 참조는 footnotes 순서대로 1..N 으로 치환된다
    assert n["sections"][0]["blocks"][1]["refs"] == [1]
    assert n["sections"][1]["blocks"][1]["refs"] == [2]
    assert n["footnotes"][0]["source"] == "1. 플랫폼 개요 · h1_intro"


def test_normalize_앵커는_선택자로_쓸_수_있어야(still):
    doc = make_doc(still)
    doc["sections"][0]["title"] = "1. 오프닝"     # 숫자로 시작하는 제목
    doc["sections"][1]["title"] = "1. 오프닝"     # 같은 제목 중복
    n = normalize_document(doc)
    ids = [s["id"] for s in n["sections"]]
    assert len(set(ids)) == 2, "중복 제목이 같은 앵커를 쓰면 안 된다"
    for sid in ids:
        assert not sid[0].isdigit(), f"CSS 선택자로 못 쓰는 앵커: {sid}"
    assert n["auto_number"] is False, "제목이 이미 번호를 달면 목차 번호를 겹치지 않는다"


def test_normalize_표_행이_매핑이어도_열_순서로_정렬(still):
    n = normalize_document(make_doc(still))
    table = n["sections"][1]["blocks"][1]
    assert table["columns"] == ["항목", "게시"]
    assert table["rows"] == [["사본 여부", "사본 아님"], ["편집 권한", "작성자 + 보직장"]]


def test_normalize_없는_그림은_건너뛰고_기록(still, tmp_path):
    doc = make_doc(still)
    doc["sections"][0]["blocks"][0]["src"] = str(tmp_path / "없는파일.png")
    n = normalize_document(doc)
    assert n["figures"] == 0
    assert len(n["missing_images"]) == 1


@pytest.mark.parametrize("mutate, msg", [
    (lambda d: d["meta"].pop("title"), "title"),
    (lambda d: d.__setitem__("sections", []), "sections"),
    (lambda d: d["sections"][0]["blocks"][1].__setitem__("refs", ["없는조각"]), "각주"),
    (lambda d: d["sections"][0]["blocks"][0].__setitem__("type", "hologram"), "type"),
])
def test_normalize_잘못된_입력은_거부(still, mutate, msg):
    doc = make_doc(still)
    mutate(doc)
    with pytest.raises(ValueError, match=msg):
        normalize_document(doc)


def test_마크다운_변환_제목_표_목록(still):
    blocks = blocks_from_markdown(
        "# 회의록\n\n결정 사항이다.\n\n- 첫째\n- 둘째\n\n| a | b |\n|---|---|\n| 1 | 2 |\n"
    )
    kinds = [b["type"] for b in blocks]
    assert kinds == ["heading", "paragraph", "list", "table"]
    assert blocks[2]["items"] == ["첫째", "둘째"]
    assert blocks[3]["columns"] == ["a", "b"] and blocks[3]["rows"] == [["1", "2"]]


# ── HTML ────────────────────────────────────────────────────────────────────

def test_export_html_light_는_상대경로_참조(still, tmp_path):
    out = tmp_path / "report.html"
    info = export_html(make_doc(still), out, mode="light", log=lambda m: None)
    html = out.read_text(encoding="utf-8")
    assert info["figures"] == 1 and info["footnotes"] == 2
    assert "data:image" not in html, "light 모드는 이미지를 인라인하지 않는다"
    assert re.search(r'<img src="(?!data:)[^"]+scene\.png"', html)
    assert "PretendardVariable.woff2" in html, "폰트는 상대 경로로 참조한다"


def test_export_html_self_는_외부_자원_0(still, tmp_path):
    out = tmp_path / "self.html"
    export_html(make_doc(still), out, mode="self", log=lambda m: None)
    html = out.read_text(encoding="utf-8")
    # 모든 src / url() 가 data: 여야 한다
    for m in re.finditer(r'(?:src="|url\(\')([^"\')]+)', html):
        assert m.group(1).startswith("data:"), f"외부 자원 참조가 남았다: {m.group(1)[:60]}"
    assert "<link" not in html and "<script" not in html


def test_export_html_목차와_각주_앵커가_모두_해소된다(still, tmp_path):
    out = tmp_path / "report.html"
    export_html(make_doc(still), out, mode="light", log=lambda m: None)
    html = out.read_text(encoding="utf-8")
    for href in re.findall(r'<a href="#([^"]+)"', html):
        assert f'id="{href}"' in html, f"앵커 대상이 없다: #{href}"
    assert html.count('class="fnref"') == 2


def test_export_html_텍스트_이스케이프(still, tmp_path):
    doc = make_doc(still)
    doc["sections"][0]["blocks"][1]["text"] = "<script>alert(1)</script> & 그 밖에"
    out = tmp_path / "esc.html"
    export_html(doc, out, mode="light", log=lambda m: None)
    html = out.read_text(encoding="utf-8")
    assert "<script>alert(1)</script>" not in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt; &amp; 그 밖에" in html


# ── PDF ─────────────────────────────────────────────────────────────────────

def test_export_pdf_A4_쪽수와_규격(still, tmp_path):
    from pypdf import PdfReader

    out = tmp_path / "report.pdf"
    info = export_pdf(make_doc(still), out, log=lambda m: None)
    assert out.exists() and info["pages"] >= 3, "표지·목차·본문이 최소 3쪽"
    reader = PdfReader(str(out))
    assert len(reader.pages) == info["pages"]
    box = reader.pages[0].mediabox
    assert abs(float(box.width) - 595.28) < 2 and abs(float(box.height) - 841.89) < 2, \
        f"A4(595×842pt)가 아니다: {box.width}×{box.height}"
    # 추출 텍스트는 자간이 공백으로 흩어지고 숫자·문장부호는 서브셋 cmap 밖으로 빠진다.
    # 공백을 지우고 한글 본문만 대조한다.
    flat = re.sub(r"\s+", "", "".join(p.extract_text() or "" for p in reader.pages))
    assert "심의문서exporter검증본" in flat              # 머리말 = 문서 제목
    assert "보고서는자기심의를각주로단다" in flat           # 요약
    assert "오프닝" in flat and "구조비교" in flat          # 절 제목·소제목
    assert re.search(r"1/\d+", flat)                    # 꼬리말 = 쪽번호(크로미엄 기본 폰트)


# ── DOCX ────────────────────────────────────────────────────────────────────

def test_export_docx_재열기_문단_표_그림(still, tmp_path):
    from docx import Document
    from docx.oxml.ns import qn

    out = tmp_path / "report.docx"
    info = export_docx(make_doc(still), out, log=lambda m: None)
    d = Document(str(out))

    assert len(d.tables) == info["tables"] == 2
    assert all(t.style.name == "Table Grid" for t in d.tables)
    assert d.tables[0].rows[0].cells[0].paragraphs[0].text == "항목"
    images = [r for r in d.part.rels.values() if "image" in r.reltype]
    assert len(images) == info["figures"] == 1
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    assert "목차" in heads and "오프닝" in heads and "부록 A. 회의록" in heads
    assert any("그림 1" in p.text for p in d.paragraphs), "그림 캡션이 없다"

    # 목차 하이퍼링크 ↔ 제목 북마크가 전부 맞물린다
    body = d.element.body
    anchors = {h.get(qn("w:anchor")) for h in body.iter(qn("w:hyperlink"))}
    bookmarks = {b.get(qn("w:name")) for b in body.iter(qn("w:bookmarkStart"))}
    assert anchors and anchors <= bookmarks
    assert info["toc_links"] == len(anchors)

    # A4 · 머리말 = 제목 · 꼬리말 = PAGE 필드
    sec = d.sections[0]
    assert round(sec.page_width.mm) == 210 and round(sec.page_height.mm) == 297
    assert sec.header.paragraphs[0].text == "심의 문서 exporter 검증본"
    footer_xml = sec.footer.part.element.xml
    assert " PAGE " in footer_xml and " NUMPAGES " in footer_xml

    # 미주(출처) — 본문 참조 표시 [1] 과 끝의 출처 목록
    assert info["endnotes"] == 2
    tail = "\n".join(p.text for p in d.paragraphs)
    assert "1. 플랫폼 개요 · h1_intro" in tail


def test_export_docx_는_편집_가능한_텍스트다(still, tmp_path):
    """그림으로 구운 문서가 아니라 문단·런이 살아 있는 진짜 워드 문서인지."""
    from docx import Document

    out = tmp_path / "edit.docx"
    export_docx(make_doc(still), out, log=lambda m: None)
    d = Document(str(out))
    body_text = "\n".join(p.text for p in d.paragraphs)
    assert "보고서는 한 번만 작성합니다." in body_text
    bolds = [r.text for p in d.paragraphs for r in p.runs if r.bold]
    assert "한 번만" in bolds, "**강조** 가 굵은 런으로 살아 있어야 한다"


def test_export_document_는_세_형식을_한_번에(still, tmp_path):
    res = export_document(make_doc(still), tmp_path, "brief", log=lambda m: None)
    assert set(res) == {"html", "pdf", "docx"}
    for ext in ("html", "pdf", "docx"):
        assert (tmp_path / f"brief.{ext}").stat().st_size > 0
    assert res["html"]["figures"] == res["pdf"]["figures"] == res["docx"]["figures"]


# ── 조립기(wdpipeline.document) 트리 수용 ──────────────────────────────────

def assembled_tree(still_path: Path) -> dict:
    """조립기가 내는 문서 트리 모양 — 절이 body/figures/tables/figure_ref/notes 로 나뉜다."""
    return {
        "meta": {"title": "조립기 트리", "core_message": "한 번만 작성한다.",
                 "date": "2026-07-29", "source_report": "ReportArchive 보고서",
                 "run_id": "delib_v2", "audience": "경영진"},
        "summary": {"lead": "요약 문단이다.", "bullets": ["첫째", "둘째"]},
        "sections": [{
            "no": 1, "heading": "게시는 사본이 아니다", "anchor": "sec-1",
            "body": ["첫 문단이다. (그림 1)", "둘째 문단이다. (표 1)"],
            "figures": [{"src": "stills/01.png", "source_path": str(still_path),
                         "caption": "오프닝 (재생 6.1초)", "no": 1, "ref": "그림 1",
                         "anchor": "fig-1"}],
            "tables": [{"caption": "비교", "columns": ["항목", "값"],
                        "rows": [["게시", "사본 아님"]], "no": 1, "ref": "표 1",
                        "anchor": "tbl-1"}],
            "figure_ref": [{"ref": "그림 1", "anchor": "fig-1", "kind": "figure", "paragraph": 0},
                           {"ref": "표 1", "anchor": "tbl-1", "kind": "table", "paragraph": 1}],
            "notes": ["[RA-001] 1. 플랫폼 개요 · h1_intro", "[RA-999] 잘린 출처 · x"],
        }],
        "appendix": {
            "deliberation": {"meeting_id": "339e3bac", "topic": "재심의 v2",
                             "participants": [{"id": "narr-copywriter", "name": "카피라이터"}],
                             "decisions": [{"text": "게시는 사본이 아니다", "turn": 12}],
                             "open_issues": [], "action_items": [], "rounds": []},
            "sources": [{"ref": "RA-001", "page": "1. 플랫폼 개요", "block_id": "h1_intro",
                         "text": "플랫폼 개요"}],
        },
        "toc": [{"no": 1, "heading": "게시는 사본이 아니다", "anchor": "sec-1"}],
    }


def test_조립기_트리를_변환없이_받는다(still):
    n = normalize_document(assembled_tree(still))
    # 요약 절이 앞에 서고, 본문 절은 조립기 anchor 를 그대로 쓴다
    assert [s["id"] for s in n["sections"]] == ["sec-summary", "sec-1"]
    kinds = [b["type"] for b in n["sections"][1]["blocks"]]
    assert kinds == ["paragraph", "figure", "paragraph", "table"], \
        "참조 토큰 자리에 그림·표가 끼워지지 않았다"
    assert n["sections"][1]["blocks"][0]["refs"] == [1]
    assert len(n["footnotes"]) == 1, "appendix.sources 밖의 note 참조는 버려야 한다"
    assert n["appendix"] and n["appendix"][0]["title"] == "부록. 심의 경과"
    assert n["meta_rows"][-1] == ("회의 ID", "339e3bac")
    assert n["figures"] == 1 and n["tables"] == 2  # 본문 1 + 심의 경과 참가자표 1


def test_조립기_실연동(tmp_path, still):
    """실제 assemble_document 산출을 그대로 세 형식으로 굽는다 (조립기가 있을 때만)."""
    pytest.importorskip("wdpipeline.document")
    from wdpipeline.document import assemble_document, validate_document

    run_src = ROOT / "data" / "pipeline" / "delib_v2"
    stills_src = ROOT / "data" / "doc_check" / "stills"
    if not (run_src / "scenario.json").exists() or not stills_src.is_dir():
        pytest.skip("delib_v2 심의 산출 또는 스틸이 없다")

    run = tmp_path / "run"
    (run / "stills").mkdir(parents=True)
    for name in ("scenario.json", "fragments.json", "report.norm.json"):
        if (run_src / name).exists():
            shutil.copy2(run_src / name, run / name)
    scen = json.loads((run / "scenario.json").read_text(encoding="utf-8"))
    cursor = 0.0
    for i, sc in enumerate(scen["scenes"], 1):
        for st in sc.get("stills") or []:
            src = stills_src / f"{i:02d}_{sc['name']}.png"
            if src.exists():
                shutil.copy2(src, run / "stills" /
                             f"{i:02d}_{sc['name']}_{cursor + float(st):.1f}s.png")
        cursor += float(sc["dur"])

    tree = assemble_document(run, style="report")
    assert validate_document(tree) == []
    info = export_html(tree, tmp_path / "a.html", mode="self", assets_root=ROOT,
                       log=lambda m: None)
    assert info["missing_images"] == [] and info["figures"] >= len(scen["scenes"])
    assert info["footnotes"] > 0, "조각 출처가 각주로 실리지 않았다"
    docx = export_docx(tree, tmp_path / "a.docx", assets_root=ROOT, log=lambda m: None)
    assert docx["figures"] == info["figures"] and docx["endnotes"] == info["footnotes"]


# ── 실데이터 재검증 ─────────────────────────────────────────────────────────

@pytest.mark.skipif(not REAL_FIXTURE.exists(), reason="data/doc_check/document.json 없음")
def test_실데이터_심의산출_3형식(tmp_path):
    """delib_v2 심의 산출(내레이션·씬 스틸·조각 구조·회의록)로 세 형식을 굽는다."""
    doc = json.loads(REAL_FIXTURE.read_text(encoding="utf-8"))
    html = export_html(doc, tmp_path / "r.html", mode="self", assets_root=ROOT,
                       log=lambda m: None)
    docx = export_docx(doc, tmp_path / "r.docx", assets_root=ROOT, log=lambda m: None)
    pdf = export_pdf(doc, tmp_path / "r.pdf", assets_root=ROOT, log=lambda m: None)
    assert html["figures"] == docx["figures"] == pdf["figures"] == 7
    assert html["tables"] == docx["tables"] == pdf["tables"]
    assert html["missing_images"] == []
    assert pdf["pages"] >= 8
    # self 모드 단일 파일은 어디로 옮겨도 열린다 — 외부 참조 0 을 텍스트로 증명
    moved = tmp_path / "moved" / "r.html"
    moved.parent.mkdir()
    shutil.copy2(tmp_path / "r.html", moved)
    text = moved.read_text(encoding="utf-8")
    assert not re.search(r'src="(?!data:)', text) and "<link" not in text
