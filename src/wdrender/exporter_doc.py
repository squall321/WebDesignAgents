# 문서 exporter — 같은 심의 산출(DocumentDoc)을 HTML(단일 파일 가능)·PDF(chromium 인쇄)·DOCX(편집 가능)로 조판한다
from __future__ import annotations

import base64
import html as _html
import json
import mimetypes
import os
import re
import tempfile
from dataclasses import asdict, is_dataclass
from pathlib import Path
from typing import Any, Callable

from .doc_theme import (
    DOCX_FONT,
    DocTheme,
    font_data_uri,
    font_face_css,
    font_relpath,
    load_doc_theme,
)

# ────────────────────────────────────────────────────────────────────────────
# DocumentDoc 계약 (P5 문서 경로 — 조립기가 만들고 이 모듈이 소비한다)
#
# {
#   "version": "1.0",
#   "meta": {
#     "title": str,               # 문서 제목 (필수)
#     "subtitle": str,            # 부제
#     "kicker": str,              # 표지 상단 라벨 (기본 "심의 보고서")
#     "core_message": str,        # 요약 상자 = 시나리오 meta.core_message
#     "audience": str, "date": str, "authors": [str],
#     "meeting_id": str|None, "source_report_id": int|None,
#     "theme": str,               # 토큰 id (기본 hwax-blue)
#   },
#   "sections": [{"id"?: str, "title": str, "level"?: 1|2, "blocks": [Block, ...]}],
#   "appendix": [{"title": str, "blocks": [Block]} | {"title": str, "markdown": str}],
#   "footnotes": [{"id": str, "text": str, "source": {"page": str, "block_id": str}}],
# }
#
# Block =
#   {"type":"paragraph", "text": str, "refs": [footnote id, ...]}   # 씬 narration
#   {"type":"heading",   "text": str}                               # 절 안의 소제목
#   {"type":"figure",    "src": path, "caption": str, "alt"?: str}  # 씬 stills 캡처
#   {"type":"table",     "columns": [str], "rows": [[str]], "caption"?: str, "refs"?: []}
#   {"type":"list",      "items": [str], "ordered"?: bool}
#   {"type":"callout",   "title"?: str, "text": str}
#   {"type":"quote",     "text": str, "by"?: str}
#
# dict·pydantic 모델·dataclass 어느 쪽으로 와도 받는다(_as_dict).
# 조립기(wdpipeline.document.assemble_document)의 문서 트리는 from_assembled 이
# 위 계약으로 옮긴다 — normalize_document 가 모양을 보고 알아서 갈아탄다.
# ────────────────────────────────────────────────────────────────────────────

BLOCK_TYPES = ("paragraph", "heading", "figure", "table", "list", "callout", "quote")

# 인라인 강조 — minutes.md 가 쓰는 **굵게** 만 해석한다(문서용 최소 문법)
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.S)


def _as_dict(obj: Any) -> dict:
    """dict / pydantic BaseModel / dataclass / JSON 문자열을 dict 로 통일한다."""
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, str):
        return json.loads(obj)
    dump = getattr(obj, "model_dump", None)  # pydantic v2
    if callable(dump):
        return dump(mode="python")
    if is_dataclass(obj) and not isinstance(obj, type):
        return asdict(obj)
    raise TypeError(f"DocumentDoc 로 해석할 수 없는 타입: {type(obj).__name__}")


def _s(v: Any) -> str:
    return "" if v is None else str(v)


def _slug(text: str, fallback: str) -> str:
    s = re.sub(r"[^0-9A-Za-z가-힣]+", "-", text).strip("-").lower()
    return s or fallback


# ── 마크다운(부록 minutes.md) → 블록 ─────────────────────────────────────────

def blocks_from_markdown(md: str) -> list[dict]:
    """회의록 수준의 최소 마크다운을 블록으로 바꾼다 — 제목·표·목록·문단만."""
    blocks: list[dict] = []
    lines = md.splitlines()
    i = 0
    para: list[str] = []
    items: list[str] = []

    def flush() -> None:
        nonlocal para, items
        if para:
            blocks.append({"type": "paragraph", "text": " ".join(para).strip()})
            para = []
        if items:
            blocks.append({"type": "list", "items": items, "ordered": False})
            items = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush()
            i += 1
            continue
        if stripped.startswith("#"):
            flush()
            text = stripped.lstrip("#").strip()
            if text:
                blocks.append({"type": "heading", "text": text})
            i += 1
            continue
        # 표 — | a | b | 다음 줄이 구분선(---)
        if stripped.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].strip().replace("|", "").replace(":", "").replace(" ", "")
        ) == {"-"}:
            flush()
            cols = [c.strip() for c in stripped.strip("|").split("|")]
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            blocks.append({"type": "table", "columns": cols, "rows": rows})
            continue
        if stripped.startswith(("- ", "* ")):
            if para:
                flush()
            items.append(stripped[2:].strip())
            i += 1
            continue
        if items:
            flush()
        para.append(stripped)
        i += 1
    flush()
    return blocks


# ── 정규화 ──────────────────────────────────────────────────────────────────

def _norm_block(raw: Any, where: str, counters: dict, fn_index: dict[str, int],
                assets_root: Path, missing: list[str]) -> dict | None:
    if not isinstance(raw, dict):
        raise ValueError(f"{where} 가 매핑이 아니다 — {{'type': ...}} 형태로 적어라")
    btype = _s(raw.get("type") or "paragraph")
    if btype not in BLOCK_TYPES:
        raise ValueError(f"{where} 의 type {btype!r} 를 모른다 — 가능한 값은 {list(BLOCK_TYPES)} 다")

    def refs() -> list[int]:
        out: list[int] = []
        for r in raw.get("refs") or []:
            n = fn_index.get(_s(r))
            if n is None:
                raise ValueError(f"{where} 의 각주 참조 {r!r} 가 footnotes 에 없다")
            out.append(n)
        return out

    if btype == "paragraph":
        text = _s(raw.get("text")).strip()
        return {"type": "paragraph", "text": text, "refs": refs()} if text else None
    if btype == "heading":
        text = _s(raw.get("text")).strip()
        return {"type": "heading", "text": text} if text else None
    if btype == "figure":
        src = _s(raw.get("src"))
        path = Path(src)
        if src and not path.is_absolute():
            path = assets_root / path
        if not src or not path.exists():
            missing.append(src or "(빈 src)")
            return None
        counters["figure"] += 1
        return {
            "type": "figure",
            "path": path,
            "caption": _s(raw.get("caption")).strip(),
            "alt": _s(raw.get("alt") or raw.get("caption")).strip(),
            "number": counters["figure"],
        }
    if btype == "table":
        # 열은 "이름" 또는 {"key","label"}(조각 structured 원형) 둘 다 받는다.
        keys: list[str] = []
        cols: list[str] = []
        for c in raw.get("columns") or []:
            if isinstance(c, dict):
                keys.append(_s(c.get("key")))
                cols.append(_s(c.get("label") or c.get("key")))
            else:
                keys.append("")
                cols.append(_s(c))
        rows: list[list[str]] = []
        for r in raw.get("rows") or []:
            if isinstance(r, dict):
                # 행이 매핑이면 열 key 순서로 정렬한다(키가 없으면 값 순서 그대로)
                rows.append([_s(r.get(k)) for k in keys] if any(keys)
                            else [_s(v) for v in r.values()])
            else:
                rows.append([_s(c) for c in r])
        if not cols and not rows:
            return None
        if not cols and rows:  # 헤더가 없으면 첫 행을 헤더로 본다
            cols, rows = rows[0], rows[1:]
        counters["table"] += 1
        return {
            "type": "table", "columns": cols, "rows": rows,
            "caption": _s(raw.get("caption")).strip(), "refs": refs(),
            "number": counters["table"],
        }
    if btype == "list":
        items = [_s(x).strip() for x in (raw.get("items") or []) if _s(x).strip()]
        return {"type": "list", "items": items, "ordered": bool(raw.get("ordered"))} if items else None
    if btype == "callout":
        text = _s(raw.get("text")).strip()
        return {"type": "callout", "title": _s(raw.get("title")).strip(), "text": text} if text else None
    # quote
    text = _s(raw.get("text")).strip()
    return {"type": "quote", "text": text, "by": _s(raw.get("by")).strip()} if text else None


def _is_assembled(d: dict) -> bool:
    """wdpipeline.document.assemble_document 이 낸 문서 트리인가 (절이 body 를 갖는다)."""
    secs = d.get("sections")
    return bool(secs) and isinstance(secs[0], dict) and "body" in secs[0]


def from_assembled(tree: Any) -> dict:
    """조립기(wdpipeline.document) 트리를 exporter 계약(DocumentDoc)으로 옮긴다.

    조립기는 절마다 본문 문단·그림·표·출처 note 를 따로 담고, 문단 안에 "(그림 1)" 같은
    참조 토큰을 심어 둔다. 여기서는 그 참조 자리(figure_ref[].paragraph)에 맞춰 그림·표를
    문단 사이에 끼워 넣어 읽는 순서대로 편다.
    """
    t = _as_dict(tree)
    meta = _as_dict(t.get("meta") or {})
    appendix_in = _as_dict(t.get("appendix") or {})
    delib = _as_dict(appendix_in.get("deliberation") or {})

    footnotes = []
    known: set[str] = set()
    for s in appendix_in.get("sources") or []:
        s = _as_dict(s)
        ref = _s(s.get("ref"))
        if not ref or ref in known:
            continue
        known.add(ref)
        footnotes.append({
            "id": ref, "text": _s(s.get("text")),
            "source": {"page": _s(s.get("page")), "block_id": _s(s.get("block_id"))},
        })

    sections: list[dict] = []
    summary = _as_dict(t.get("summary") or {})
    if summary.get("lead") or summary.get("bullets"):
        blocks: list[dict] = []
        if summary.get("lead"):
            blocks.append({"type": "paragraph", "text": _s(summary["lead"])})
        if summary.get("bullets"):
            blocks.append({"type": "list", "items": [_s(b) for b in summary["bullets"]]})
        sections.append({"id": "sec-summary", "title": "요약", "blocks": blocks})

    for raw in t.get("sections") or []:
        sec = _as_dict(raw)
        by_ref = {}
        for fig in sec.get("figures") or []:
            f = _as_dict(fig)
            by_ref[_s(f.get("ref"))] = {
                "type": "figure",
                "src": _s(f.get("source_path") or f.get("src")),
                "caption": _s(f.get("caption")),
                "alt": _s(f.get("caption")),
            }
        for tbl in sec.get("tables") or []:
            b = _as_dict(tbl)
            by_ref[_s(b.get("ref"))] = {
                "type": "table", "caption": _s(b.get("caption")),
                "columns": b.get("columns") or [], "rows": b.get("rows") or [],
            }
        # 문단 슬롯 → 그 뒤에 놓을 참조 블록
        slots: dict[int, list[dict]] = {}
        for r in sec.get("figure_ref") or []:
            r = _as_dict(r)
            block = by_ref.pop(_s(r.get("ref")), None)
            if block:
                slots.setdefault(int(r.get("paragraph") or 0), []).append(block)
        # 절의 출처 note 는 첫 문단에 단다 (조립기가 문단 단위로는 기록하지 않는다)
        refs = [m.group(1) for m in
                (re.match(r"^\[([^\]]+)\]", _s(n)) for n in sec.get("notes") or []) if m]
        refs = [r for r in refs if r in known]

        blocks = []
        paras = [p for p in (sec.get("body") or []) if _s(p).strip()]
        for i, para in enumerate(paras):
            blocks.append({"type": "paragraph", "text": _s(para),
                           "refs": refs if i == 0 else []})
            blocks.extend(slots.pop(i, []))
        for extra in slots.values():         # 슬롯 밖으로 밀린 참조는 절 끝에
            blocks.extend(extra)
        blocks.extend(by_ref.values())       # 참조 토큰이 없는 그림·표도 버리지 않는다
        sections.append({"id": _s(sec.get("anchor")) or "", "title": _s(sec.get("heading")),
                         "blocks": blocks})

    appendix = _appendix_from_deliberation(delib)
    return {
        "version": "1.0",
        "meta": {
            "title": _s(meta.get("title")),
            "subtitle": _s(meta.get("source_report")),
            "core_message": _s(meta.get("core_message")),
            "audience": _s(meta.get("audience")),
            "date": _s(meta.get("date")),
            "authors": [_s(_as_dict(p).get("name")) for p in delib.get("participants") or []],
            "meeting_id": _s(delib.get("meeting_id")) or None,
        },
        "sections": sections,
        "appendix": appendix,
        "footnotes": footnotes,
    }


def _appendix_from_deliberation(delib: dict) -> list[dict]:
    """심의 경과(참가자·결정·미해결 쟁점·액션아이템·라운드)를 부록 한 절로 편다."""
    if not delib:
        return []
    blocks: list[dict] = []
    if delib.get("topic"):
        blocks.append({"type": "paragraph", "text": _s(delib["topic"])})
    parts = [_as_dict(p) for p in delib.get("participants") or []]
    if parts:
        blocks.append({"type": "heading", "text": "참가자"})
        blocks.append({"type": "table", "columns": ["ID", "이름"],
                       "rows": [[_s(p.get("id")), _s(p.get("name"))] for p in parts]})
    for key, title in (("decisions", "결정"), ("open_issues", "미해결 쟁점")):
        items = [_as_dict(x) for x in delib.get(key) or []]
        if items:
            blocks.append({"type": "heading", "text": title})
            blocks.append({"type": "list", "items": [
                _s(x.get("text")) + (f" (턴 #{x['turn']})" if x.get("turn") else "")
                for x in items]})
    actions = [_as_dict(x) for x in delib.get("action_items") or []]
    if actions:
        blocks.append({"type": "heading", "text": "액션아이템"})
        blocks.append({"type": "table", "columns": ["#", "내용", "담당", "턴"],
                       "rows": [[_s(a.get("no")), _s(a.get("text")), _s(a.get("owner")),
                                 _s(a.get("turn"))] for a in actions]})
    rounds = [_as_dict(x) for x in delib.get("rounds") or []]
    if rounds:
        blocks.append({"type": "heading", "text": "라운드"})
        blocks.append({"type": "table", "columns": ["라운드", "성격", "턴", "발언자"],
                       "rows": [[_s(r.get("round")), _s(r.get("label")), _s(r.get("turns")),
                                 ", ".join(_s(s) for s in r.get("speakers") or [])]
                                for r in rounds]})
    return [{"id": "apx-deliberation", "title": "부록. 심의 경과", "blocks": blocks}] if blocks else []


def normalize_document(doc: Any, *, assets_root: Path | None = None) -> dict:
    """DocumentDoc(또는 조립기 트리)을 조판기가 바로 쓰는 모양으로 정규화한다.

    - 절마다 앵커 id 를 확정(중복은 -2, -3 … 접미)
    - 그림·표에 일련번호를 매긴다
    - 각주는 footnotes 목록 순서대로 1..N, 블록의 refs 는 그 번호로 치환
    - 부록의 markdown 은 블록으로 변환
    - 존재하지 않는 그림 파일은 건너뛰고 missing_images 에 기록
    """
    d = _as_dict(doc)
    if _is_assembled(d):  # 조립기 트리는 계약 모양으로 옮겨서 받는다
        d = from_assembled(d)
    root = Path(assets_root) if assets_root else Path.cwd()
    meta = _as_dict(d.get("meta") or {})
    title = _s(meta.get("title")).strip()
    if not title:
        raise ValueError("meta.title 이 비었다 — 문서 제목은 필수다")

    footnotes_raw = [_as_dict(f) for f in (d.get("footnotes") or [])]
    fn_index: dict[str, int] = {}
    footnotes: list[dict] = []
    for n, f in enumerate(footnotes_raw, 1):
        fid = _s(f.get("id")) or f"fn-{n}"
        fn_index[fid] = n
        src = _as_dict(f.get("source") or {})
        bits = [b for b in (_s(src.get("page")), _s(src.get("block_id"))) if b]
        footnotes.append({
            "n": n, "id": fid, "text": _s(f.get("text")).strip(),
            "source": " · ".join(bits),
        })

    counters = {"figure": 0, "table": 0}
    missing: list[str] = []
    used_ids: set[str] = set()

    def norm_group(raw_list: Any, kind: str) -> list[dict]:
        out: list[dict] = []
        prefix = "sec" if kind == "sections" else "apx"
        for i, raw in enumerate(raw_list or [], 1):
            s = _as_dict(raw)
            stitle = _s(s.get("title")).strip()
            sid = _s(s.get("id")).strip() or f"{prefix}-{_slug(stitle, str(i))}"
            # 숫자로 시작하는 id 는 CSS 선택자로 못 쓴다("1. 오프닝" 같은 제목) — 접두를 붙인다
            if sid[0].isdigit() or sid[0] == "-":
                sid = f"{prefix}-{sid}"
            base, k = sid, 2
            while sid in used_ids:
                sid, k = f"{base}-{k}", k + 1
            used_ids.add(sid)
            blocks_raw = s.get("blocks")
            if blocks_raw is None and s.get("markdown") is not None:
                blocks_raw = blocks_from_markdown(_s(s.get("markdown")))
            blocks: list[dict] = []
            for j, b in enumerate(blocks_raw or []):
                nb = _norm_block(b, f"{kind}[{i - 1}].blocks[{j}]", counters,
                                 fn_index, root, missing)
                if nb:
                    blocks.append(nb)
            level = int(s.get("level") or 1)
            out.append({"id": sid, "title": stitle, "level": max(1, min(2, level)),
                        "blocks": blocks})
        return out

    sections = norm_group(d.get("sections"), "sections")
    if not sections:
        raise ValueError("sections 가 비었다 — 절을 1개 이상 넣어라")
    appendix = norm_group(d.get("appendix"), "appendix")

    meta_rows = [
        (label, _s(meta.get(key)).strip())
        for label, key in (("작성일", "date"), ("대상", "audience"))
    ]
    authors = [_s(a) for a in (meta.get("authors") or []) if _s(a)]
    if authors:
        meta_rows.append(("심의", ", ".join(authors)))
    if meta.get("meeting_id"):
        meta_rows.append(("회의 ID", _s(meta["meeting_id"])))
    if meta.get("source_report_id") is not None:
        meta_rows.append(("원천 보고서", _s(meta["source_report_id"])))

    # 절 제목이 이미 "1. …" 처럼 번호를 달고 있으면 목차 번호를 겹쳐 매기지 않는다
    auto_number = not any(
        re.match(r"^\s*\d+[.)]", s["title"]) for s in sections if s["title"]
    )
    return {
        "title": title,
        "subtitle": _s(meta.get("subtitle")).strip(),
        "auto_number": auto_number,
        "kicker": _s(meta.get("kicker")).strip() or "심의 보고서",
        "core_message": _s(meta.get("core_message")).strip(),
        "theme": _s(meta.get("theme")).strip() or "hwax-blue",
        "meta_rows": [(k, v) for k, v in meta_rows if v],
        "sections": sections,
        "appendix": appendix,
        "footnotes": footnotes,
        "figures": counters["figure"],
        "tables": counters["table"],
        "missing_images": missing,
    }


# ── HTML ────────────────────────────────────────────────────────────────────

def _inline_runs(text: str) -> list[tuple[str, bool]]:
    """**굵게** 만 해석해 (텍스트, bold) 런으로 쪼갠다."""
    runs: list[tuple[str, bool]] = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False))
        runs.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False))
    return runs or [(text, False)]


def _inline_html(text: str) -> str:
    return "".join(
        f"<strong>{_html.escape(t)}</strong>" if b else _html.escape(t)
        for t, b in _inline_runs(text)
    )


def _refs_html(refs: list[int]) -> str:
    if not refs:
        return ""
    links = "".join(
        f'<a href="#fn-{n}" id="fnref-{n}">{n}</a>' for n in refs
    )
    return f'<sup class="fnref">{links}</sup>'


def _img_src(path: Path, out_dir: Path, embed: bool) -> str:
    if embed:
        mime = mimetypes.guess_type(path.name)[0] or "image/png"
        b64 = base64.b64encode(path.read_bytes()).decode("ascii")
        return f"data:{mime};base64,{b64}"
    return os.path.relpath(path, out_dir).replace("\\", "/")


def _blocks_html(blocks: list[dict], out_dir: Path, embed: bool) -> str:
    parts: list[str] = []
    for b in blocks:
        t = b["type"]
        if t == "paragraph":
            parts.append(f"<p>{_inline_html(b['text'])}{_refs_html(b['refs'])}</p>")
        elif t == "heading":
            parts.append(f"<h3>{_inline_html(b['text'])}</h3>")
        elif t == "figure":
            src = _img_src(b["path"], out_dir, embed)
            cap = (
                f'<figcaption><span class="num">그림 {b["number"]}</span>'
                f"{_inline_html(b['caption'])}</figcaption>" if b["caption"]
                else f'<figcaption><span class="num">그림 {b["number"]}</span></figcaption>'
            )
            parts.append(
                f'<figure><img src="{src}" alt="{_html.escape(b["alt"])}">{cap}</figure>'
            )
        elif t == "table":
            head = "".join(f"<th>{_inline_html(c)}</th>" for c in b["columns"])
            body = "".join(
                "<tr>" + "".join(f"<td>{_inline_html(c)}</td>" for c in r) + "</tr>"
                for r in b["rows"]
            )
            cap = (
                f'<caption><span class="num">표 {b["number"]}</span>'
                f"{_inline_html(b['caption'])}{_refs_html(b['refs'])}</caption>"
                if b["caption"] or b["refs"]
                else f'<caption><span class="num">표 {b["number"]}</span></caption>'
            )
            parts.append(
                f'<div class="table-wrap"><table>{cap}<thead><tr>{head}</tr></thead>'
                f"<tbody>{body}</tbody></table></div>"
            )
        elif t == "list":
            tag = "ol" if b["ordered"] else "ul"
            lis = "".join(f"<li>{_inline_html(x)}</li>" for x in b["items"])
            parts.append(f"<{tag}>{lis}</{tag}>")
        elif t == "callout":
            head = f'<b>{_inline_html(b["title"])}</b>' if b["title"] else ""
            parts.append(f'<aside class="callout">{head}<p>{_inline_html(b["text"])}</p></aside>')
        elif t == "quote":
            by = f'<cite>{_inline_html(b["by"])}</cite>' if b["by"] else ""
            parts.append(f"<blockquote><p>{_inline_html(b['text'])}</p>{by}</blockquote>")
    return "\n".join(parts)


def _css(theme: DocTheme, font_src: str | None) -> str:
    face = font_face_css(font_src) if font_src else ""
    return f"""{face}
*{{box-sizing:border-box;}}
html{{-webkit-text-size-adjust:100%;}}
body{{margin:0;background:{theme.shade};color:{theme.ink};
  font-family:{theme.font_stack};font-size:16px;line-height:1.7;
  font-feature-settings:'tnum' 1;}}
.doc{{max-width:760px;margin:0 auto;padding:56px 28px 96px;background:{theme.page};}}
.cover{{padding:12px 0 40px;border-bottom:1px solid {theme.line};margin-bottom:40px;}}
.kicker{{display:flex;align-items:center;gap:12px;color:{theme.accent};
  font-size:13px;font-weight:700;letter-spacing:.14em;text-transform:uppercase;}}
.kicker::before{{content:'';width:34px;height:4px;background:{theme.accent};border-radius:2px;}}
h1{{font-size:38px;line-height:1.28;margin:20px 0 0;letter-spacing:-.02em;}}
.subtitle{{margin:14px 0 0;font-size:18px;color:{theme.sub};line-height:1.6;}}
.meta{{display:grid;grid-template-columns:auto 1fr;gap:6px 18px;margin:28px 0 0;
  font-size:13.5px;color:{theme.sub};}}
.meta dt{{color:{theme.faint};}}
.meta dd{{margin:0;}}
.summary{{background:{theme.accent_soft};border:1px solid {theme.accent_border};
  border-radius:14px;padding:20px 24px;margin:0 0 40px;}}
.summary .label{{font-size:12.5px;font-weight:700;letter-spacing:.1em;color:{theme.accent};
  text-transform:uppercase;}}
.summary p{{margin:8px 0 0;font-size:18px;line-height:1.6;color:{theme.ink};font-weight:600;}}
nav.toc{{border:1px solid {theme.line};border-radius:14px;padding:20px 24px;margin:0 0 44px;}}
nav.toc .label{{font-size:12.5px;font-weight:700;letter-spacing:.1em;color:{theme.faint};
  text-transform:uppercase;}}
nav.toc ol{{margin:10px 0 0;padding:0;list-style:none;counter-reset:toc;}}
nav.toc li{{counter-increment:toc;margin:2px 0;}}
nav.toc li.sub{{padding-left:22px;}}
nav.toc a{{display:flex;gap:12px;color:{theme.ink};text-decoration:none;padding:3px 0;}}
nav.toc a::before{{content:counter(toc,decimal-leading-zero);color:{theme.accent};
  font-weight:700;font-size:13px;min-width:24px;}}
nav.toc ol.plain a::before{{content:none;}}
nav.toc a:hover{{color:{theme.accent};}}
section.sec{{margin:0 0 40px;}}
h2{{font-size:25px;line-height:1.35;margin:44px 0 4px;letter-spacing:-.01em;
  padding-top:14px;border-top:2px solid {theme.accent};}}
h2.lv2{{font-size:21px;border-top:1px solid {theme.line};}}
h3{{font-size:17.5px;margin:28px 0 -6px;color:{theme.accent2};}}
p{{margin:16px 0;}}
ul,ol{{margin:16px 0;padding-left:22px;}}
li{{margin:6px 0;}}
figure{{margin:28px 0;text-align:center;}}
figure img{{max-width:100%;height:auto;border:1px solid {theme.line};border-radius:10px;
  display:block;margin:0 auto;}}
figcaption{{margin-top:10px;font-size:13.5px;color:{theme.sub};line-height:1.5;}}
.num{{color:{theme.accent};font-weight:700;margin-right:8px;}}
.table-wrap{{overflow-x:auto;margin:28px 0;}}
table{{border-collapse:collapse;width:100%;font-size:14.5px;}}
caption{{caption-side:top;text-align:left;font-size:13.5px;color:{theme.sub};
  padding-bottom:10px;}}
th,td{{text-align:left;padding:9px 12px;border-bottom:1px solid {theme.line};
  vertical-align:top;line-height:1.55;word-break:keep-all;overflow-wrap:break-word;}}
thead th{{border-bottom:1.5px solid {theme.ink};font-weight:700;font-size:13.5px;
  color:{theme.ink};}}
tbody tr:last-child td{{border-bottom:1px solid {theme.line};}}
.callout{{border-left:3px solid {theme.accent};background:{theme.shade};
  padding:14px 18px;margin:24px 0;border-radius:0 10px 10px 0;}}
.callout b{{color:{theme.accent};font-size:14px;}}
.callout p{{margin:4px 0 0;font-size:15px;}}
blockquote{{margin:24px 0;padding:0 0 0 20px;border-left:3px solid {theme.accent_border};
  color:{theme.sub};}}
blockquote cite{{display:block;margin-top:8px;font-size:13.5px;color:{theme.faint};
  font-style:normal;}}
sup.fnref a{{color:{theme.accent};text-decoration:none;font-size:11.5px;font-weight:700;
  padding:0 1px;}}
sup.fnref a::before{{content:'[';}}
sup.fnref a::after{{content:']';}}
.notes{{margin-top:56px;padding-top:20px;border-top:1px solid {theme.line};
  font-size:13.5px;color:{theme.sub};}}
.notes h2{{font-size:16px;border:0;padding:0;margin:0 0 10px;color:{theme.faint};
  letter-spacing:.08em;text-transform:uppercase;}}
.notes ol{{padding-left:20px;}}
.notes li{{margin:6px 0;line-height:1.6;}}
.notes .src{{color:{theme.faint};}}
.notes a{{color:{theme.accent};text-decoration:none;}}
@media (max-width:640px){{
  .doc{{padding:32px 18px 64px;}}
  h1{{font-size:29px;}} h2{{font-size:21px;}} body{{font-size:15.5px;}}
}}
@media print{{
  @page{{size:A4;margin:18mm 16mm;}}
  body{{background:#fff;font-size:10.5pt;}}
  .doc{{max-width:none;padding:0;}}
  .cover{{break-after:page;border-bottom:0;}}
  nav.toc{{break-after:page;}}
  nav.toc a{{color:{theme.ink};}}
  h2{{break-after:avoid;break-inside:avoid;}}
  h3{{break-after:avoid;}}
  figure,table,tr,.callout,blockquote{{break-inside:avoid;}}
  .table-wrap{{overflow:visible;}}
  figure img{{max-height:150mm;object-fit:contain;}}
  p{{orphans:2;widows:2;}}
}}"""


def render_html(ndoc: dict, *, out_dir: Path, mode: str = "light",
                theme: DocTheme | None = None) -> str:
    """정규화된 문서를 완결된 HTML 문자열로 조판한다.

    mode="self" 면 CSS·폰트·이미지를 전부 인라인해 **단일 파일**로 만든다.
    """
    if mode not in ("light", "self"):
        raise ValueError(f"mode 는 'light' 또는 'self' — 받은 값 {mode!r}")
    th = theme or load_doc_theme(ndoc["theme"])
    embed = mode == "self"
    font_src = font_data_uri() if embed else font_relpath(out_dir)

    toc = "".join(
        f'<li class="{"sub" if s["level"] == 2 else "main"}">'
        f'<a href="#{s["id"]}">{_html.escape(s["title"])}</a></li>'
        for s in ndoc["sections"] + ndoc["appendix"] if s["title"]
    )
    body: list[str] = []
    for s in ndoc["sections"] + ndoc["appendix"]:
        cls = "sec" + (" lv2" if s["level"] == 2 else "")
        head = (
            f'<h2 class="{"lv2" if s["level"] == 2 else ""}" id="{s["id"]}">'
            f"{_html.escape(s['title'])}</h2>" if s["title"] else ""
        )
        body.append(
            f'<section class="{cls}">{head}\n'
            f'{_blocks_html(s["blocks"], out_dir, embed)}</section>'
        )

    notes = ""
    if ndoc["footnotes"]:
        lis = "".join(
            f'<li id="fn-{f["n"]}">{_inline_html(f["text"])}'
            + (f' <span class="src">— {_html.escape(f["source"])}</span>' if f["source"] else "")
            + f' <a href="#fnref-{f["n"]}">↩</a></li>'
            for f in ndoc["footnotes"]
        )
        notes = f'<section class="notes"><h2>출처</h2><ol>{lis}</ol></section>'

    meta_html = ""
    if ndoc["meta_rows"]:
        meta_html = '<dl class="meta">' + "".join(
            f"<dt>{_html.escape(k)}</dt><dd>{_html.escape(v)}</dd>"
            for k, v in ndoc["meta_rows"]
        ) + "</dl>"
    summary = (
        f'<div class="summary"><div class="label">핵심 메시지</div>'
        f'<p>{_inline_html(ndoc["core_message"])}</p></div>'
        if ndoc["core_message"] else ""
    )
    subtitle = (
        f'<p class="subtitle">{_html.escape(ndoc["subtitle"])}</p>'
        if ndoc["subtitle"] else ""
    )
    return f"""<!doctype html>
<html lang="ko">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_html.escape(ndoc["title"])}</title>
<style>{_css(th, font_src)}</style>
</head>
<body>
<article class="doc">
<header class="cover">
<div class="kicker">{_html.escape(ndoc["kicker"])}</div>
<h1>{_html.escape(ndoc["title"])}</h1>
{subtitle}
{meta_html}
</header>
{summary}
<nav class="toc"><div class="label">목차</div>
<ol class="{'' if ndoc['auto_number'] else 'plain'}">{toc}</ol></nav>
{chr(10).join(body)}
{notes}
</article>
</body>
</html>"""


def export_html(doc: Any, out_path: str | Path, *, mode: str = "light",
                assets_root: Path | None = None, theme: DocTheme | None = None,
                log: Callable[[str], None] = print) -> dict:
    """DocumentDoc → HTML 파일.

    mode="light" — CSS 는 인라인, 폰트·이미지는 상대 경로 참조(원본 트리 안에서 본다).
    mode="self"  — 폰트·이미지까지 base64 로 넣은 **단일 파일**(메일 첨부·채팅 임베드용).
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ndoc = normalize_document(doc, assets_root=assets_root)
    html = render_html(ndoc, out_dir=out_path.parent, mode=mode, theme=theme)
    out_path.write_text(html, encoding="utf-8")
    size = out_path.stat().st_size
    log(f"[export_html] {out_path} ({size:,} bytes, mode={mode}, "
        f"절 {len(ndoc['sections'])} 그림 {ndoc['figures']} 표 {ndoc['tables']})")
    return {
        "out": str(out_path), "bytes": size, "mode": mode,
        "sections": len(ndoc["sections"]), "appendix": len(ndoc["appendix"]),
        "figures": ndoc["figures"], "tables": ndoc["tables"],
        "footnotes": len(ndoc["footnotes"]), "missing_images": ndoc["missing_images"],
    }


# ── PDF (chromium print-to-pdf) ─────────────────────────────────────────────

_CHROME_BAR = (
    "font-family:sans-serif;font-size:8px;width:100%;padding:0 16mm;"
    "color:#57607A;-webkit-print-color-adjust:exact;"
)


def export_pdf(doc: Any, out_path: str | Path, *, assets_root: Path | None = None,
               theme: DocTheme | None = None, margin: dict[str, str] | None = None,
               log: Callable[[str], None] = print) -> dict:
    """DocumentDoc → A4 PDF. self 모드 HTML 을 임시로 굽고 chromium 으로 인쇄한다.

    머리말=문서 제목, 꼬리말=쪽번호. 그림·표는 CSS break-inside:avoid 로 쪼개지지 않는다.
    """
    from playwright.sync_api import sync_playwright

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ndoc = normalize_document(doc, assets_root=assets_root)

    with tempfile.TemporaryDirectory(prefix="wda-doc-pdf-") as tmp:
        tmp_html = Path(tmp) / "print.html"
        # self 모드로 구우면 이미지·폰트가 전부 인라인이라 file:// 단독으로 완전히 렌더된다
        tmp_html.write_text(
            render_html(ndoc, out_dir=tmp_html.parent, mode="self", theme=theme),
            encoding="utf-8",
        )
        title = _html.escape(ndoc["title"])
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(tmp_html.as_uri(), wait_until="load")
            page.evaluate("() => document.fonts.ready.then(() => true)")
            page.emulate_media(media="print")
            page.pdf(
                path=str(out_path),
                format="A4",
                print_background=True,
                display_header_footer=True,
                header_template=f'<div style="{_CHROME_BAR}text-align:left;">{title}</div>',
                footer_template=(
                    f'<div style="{_CHROME_BAR}text-align:center;">'
                    '<span class="pageNumber"></span> / <span class="totalPages"></span></div>'
                ),
                margin=margin or {"top": "20mm", "bottom": "18mm",
                                  "left": "16mm", "right": "16mm"},
            )
            browser.close()

    size = out_path.stat().st_size
    pages = _pdf_page_count(out_path)
    log(f"[export_pdf] {out_path} ({size:,} bytes, {pages}쪽)")
    return {
        "out": str(out_path), "bytes": size, "pages": pages,
        "figures": ndoc["figures"], "tables": ndoc["tables"],
        "missing_images": ndoc["missing_images"],
    }


def _pdf_page_count(path: Path) -> int:
    """PDF 쪽수 — pypdf 가 있으면 그것으로, 없으면 /Type /Page 계수로."""
    try:
        from pypdf import PdfReader

        return len(PdfReader(str(path)).pages)
    except ImportError:
        blob = path.read_bytes()
        return blob.count(b"/Type /Page") - blob.count(b"/Type /Pages")


# ── DOCX (편집 가능한 워드 문서) ────────────────────────────────────────────

_A4_W_MM, _A4_H_MM = 210, 297
_DOCX_MARGIN_MM = 22


def _docx_set_font(run, name: str = DOCX_FONT) -> None:
    """latin 과 동아시아 타입페이스를 함께 지정 — 한글이 대체 폰트로 흘러가지 않게."""
    from docx.oxml.ns import qn

    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _docx_color(run, hex_color: str) -> None:
    from docx.shared import RGBColor

    s = hex_color.lstrip("#")
    run.font.color.rgb = RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


def _docx_bookmark(paragraph, name: str, bid: int) -> None:
    from docx.oxml.ns import qn

    p = paragraph._p
    start = p.makeelement(qn("w:bookmarkStart"), {qn("w:id"): str(bid), qn("w:name"): name})
    end = p.makeelement(qn("w:bookmarkEnd"), {qn("w:id"): str(bid)})
    p.insert(0, start)
    p.append(end)


def _docx_link(paragraph, anchor: str, text: str, theme: DocTheme):
    """문서 내부 앵커로 가는 하이퍼링크 런을 만든다(목차 항목용)."""
    from docx.oxml.ns import qn
    from docx.text.run import Run

    link = paragraph._p.makeelement(qn("w:hyperlink"), {qn("w:anchor"): anchor})
    paragraph._p.append(link)
    r = link.makeelement(qn("w:r"), {})
    link.append(r)
    run = Run(r, paragraph)
    run.text = text
    _docx_set_font(run)
    _docx_color(run, theme.ink)
    return run


def _docx_field(paragraph, instr: str, placeholder: str = "1"):
    """PAGE / NUMPAGES 같은 단순 필드를 넣는다(워드가 열 때 값을 채운다)."""
    from docx.oxml.ns import qn

    fld = paragraph._p.makeelement(qn("w:fldSimple"), {qn("w:instr"): instr})
    r = fld.makeelement(qn("w:r"), {})
    t = r.makeelement(qn("w:t"), {})
    t.text = placeholder
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _docx_runs(paragraph, text: str, theme: DocTheme, *, size_pt: float | None = None,
               color: str | None = None, italic: bool = False) -> None:
    from docx.shared import Pt

    for chunk, bold in _inline_runs(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = italic
        if size_pt:
            run.font.size = Pt(size_pt)
        _docx_set_font(run)
        _docx_color(run, color or theme.ink)


def _docx_endnote_marks(paragraph, refs: list[int], theme: DocTheme) -> None:
    from docx.shared import Pt

    for n in refs:
        run = paragraph.add_run(f"[{n}]")
        run.font.superscript = True
        run.font.size = Pt(8)
        run.bold = True
        _docx_set_font(run)
        _docx_color(run, theme.accent)


def _docx_blocks(document, blocks: list[dict], theme: DocTheme, content_w) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for b in blocks:
        t = b["type"]
        if t == "paragraph":
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            _docx_runs(p, b["text"], theme)
            _docx_endnote_marks(p, b["refs"], theme)
        elif t == "heading":
            h = document.add_heading(level=3)
            _docx_runs(h, b["text"], theme, size_pt=13, color=theme.accent2)
        elif t == "figure":
            document.add_picture(str(b["path"]), width=content_w)
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = document.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(14)
            lead = cap.add_run(f"그림 {b['number']}  ")
            lead.bold = True
            lead.font.size = Pt(9)
            _docx_set_font(lead)
            _docx_color(lead, theme.accent)
            if b["caption"]:
                _docx_runs(cap, b["caption"], theme, size_pt=9, color=theme.sub)
        elif t == "table":
            cap = document.add_paragraph()
            cap.paragraph_format.space_after = Pt(4)
            lead = cap.add_run(f"표 {b['number']}  ")
            lead.bold = True
            lead.font.size = Pt(9)
            _docx_set_font(lead)
            _docx_color(lead, theme.accent)
            if b["caption"]:
                _docx_runs(cap, b["caption"], theme, size_pt=9, color=theme.sub)
            _docx_endnote_marks(cap, b["refs"], theme)
            cols = b["columns"] or (b["rows"][0] if b["rows"] else [""])
            table = document.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, c in enumerate(cols):
                cell_p = table.rows[0].cells[i].paragraphs[0]
                _docx_runs(cell_p, c, theme, size_pt=9.5)
                for r in cell_p.runs:
                    r.bold = True
            for row in b["rows"]:
                cells = table.add_row().cells
                for i, c in enumerate(row[: len(cols)]):
                    _docx_runs(cells[i].paragraphs[0], c, theme, size_pt=9.5)
            document.add_paragraph().paragraph_format.space_after = Pt(6)
        elif t == "list":
            style = "List Number" if b["ordered"] else "List Bullet"
            for item in b["items"]:
                p = document.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(2)
                _docx_runs(p, item, theme)
        elif t == "callout":
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_after = Pt(10)
            if b["title"]:
                _docx_runs(p, f"{b['title']}  ", theme, size_pt=10.5, color=theme.accent)
                p.runs[-1].bold = True
            _docx_runs(p, b["text"], theme, size_pt=10.5, color=theme.sub)
        elif t == "quote":
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _docx_runs(p, b["text"], theme, italic=True, color=theme.sub)
            if b["by"]:
                q = document.add_paragraph()
                q.paragraph_format.left_indent = Pt(18)
                _docx_runs(q, f"— {b['by']}", theme, size_pt=9, color=theme.faint)


def export_docx(doc: Any, out_path: str | Path, *, assets_root: Path | None = None,
                theme: DocTheme | None = None,
                log: Callable[[str], None] = print) -> dict:
    """DocumentDoc → 편집 가능한 DOCX.

    제목 스타일·목차(내부 앵커 하이퍼링크)·그림(캡션 포함)·표(격자)·미주(출처)를 갖춘
    진짜 워드 문서다. 각주 API 가 python-docx 에 없으므로 출처는 문서 끝 미주로 모은다.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ndoc = normalize_document(doc, assets_root=assets_root)
    th = theme or load_doc_theme(ndoc["theme"])

    document = Document()
    sec = document.sections[0]
    sec.page_width, sec.page_height = Mm(_A4_W_MM), Mm(_A4_H_MM)
    sec.left_margin = sec.right_margin = Mm(_DOCX_MARGIN_MM)
    sec.top_margin = sec.bottom_margin = Mm(_DOCX_MARGIN_MM)
    content_w = Mm(_A4_W_MM - 2 * _DOCX_MARGIN_MM)

    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = DOCX_FONT
    normal.paragraph_format.line_spacing = 1.45

    # 머리말(문서 제목) · 꼬리말(쪽번호)
    head_p = sec.header.paragraphs[0]
    _docx_runs(head_p, ndoc["title"], th, size_pt=8.5, color=th.faint)
    foot_p = sec.footer.paragraphs[0]
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_field(foot_p, " PAGE ")
    _docx_runs(foot_p, " / ", th, size_pt=8.5, color=th.faint)
    _docx_field(foot_p, " NUMPAGES ")
    for r in foot_p.runs:
        r.font.size = Pt(8.5)
        _docx_set_font(r)
        _docx_color(r, th.faint)

    # 표지
    kick = document.add_paragraph()
    _docx_runs(kick, ndoc["kicker"], th, size_pt=9, color=th.accent)
    for r in kick.runs:
        r.bold = True
    title_p = document.add_paragraph(style="Title")
    _docx_runs(title_p, ndoc["title"], th, size_pt=26)
    if ndoc["subtitle"]:
        sub_p = document.add_paragraph()
        _docx_runs(sub_p, ndoc["subtitle"], th, size_pt=13, color=th.sub)
    for label, value in ndoc["meta_rows"]:
        m = document.add_paragraph()
        m.paragraph_format.space_after = Pt(0)
        _docx_runs(m, f"{label}  ", th, size_pt=9, color=th.faint)
        _docx_runs(m, value, th, size_pt=9, color=th.sub)

    if ndoc["core_message"]:
        document.add_paragraph()
        s = document.add_paragraph()
        _docx_runs(s, "핵심 메시지  ", th, size_pt=9, color=th.accent)
        s.runs[-1].bold = True
        core = document.add_paragraph()
        core.paragraph_format.left_indent = Pt(14)
        _docx_runs(core, ndoc["core_message"], th, size_pt=12)
        for r in core.runs:
            r.bold = True

    # 목차 — 절 제목마다 북마크를 걸고 내부 하이퍼링크로 잇는다
    groups = ndoc["sections"] + ndoc["appendix"]
    toc_head = document.add_heading(level=1)
    _docx_runs(toc_head, "목차", th, size_pt=17, color=th.ink)
    n_links = 0
    for i, s in enumerate(groups, 1):
        if not s["title"]:
            continue
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(14 if s["level"] == 2 else 0)
        if ndoc["auto_number"]:
            num = p.add_run(f"{i:02d}   ")
            num.bold = True
            num.font.size = Pt(9.5)
            _docx_set_font(num)
            _docx_color(num, th.accent)
        _docx_link(p, f"bm-{s['id']}", s["title"], th)
        n_links += 1
    document.add_page_break()

    # 본문 + 부록
    for i, s in enumerate(groups, 1):
        if s["title"]:
            h = document.add_heading(level=1 if s["level"] == 1 else 2)
            _docx_runs(h, s["title"], th, size_pt=17 if s["level"] == 1 else 14)
            _docx_bookmark(h, f"bm-{s['id']}", i)
        _docx_blocks(document, s["blocks"], th, content_w)

    # 미주 — 조각 출처(page/block_id)
    if ndoc["footnotes"]:
        h = document.add_heading(level=1)
        _docx_runs(h, "출처", th, size_pt=15, color=th.faint)
        for f in ndoc["footnotes"]:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            mark = p.add_run(f"[{f['n']}] ")
            mark.bold = True
            mark.font.size = Pt(9)
            _docx_set_font(mark)
            _docx_color(mark, th.accent)
            _docx_runs(p, f["text"], th, size_pt=9, color=th.sub)
            if f["source"]:
                _docx_runs(p, f"  — {f['source']}", th, size_pt=9, color=th.faint)

    document.save(out_path)
    size = out_path.stat().st_size
    log(f"[export_docx] {out_path} ({size:,} bytes, 절 {len(groups)} "
        f"그림 {ndoc['figures']} 표 {ndoc['tables']} 미주 {len(ndoc['footnotes'])})")
    return {
        "out": str(out_path), "bytes": size, "sections": len(ndoc["sections"]),
        "appendix": len(ndoc["appendix"]), "figures": ndoc["figures"],
        "tables": ndoc["tables"], "endnotes": len(ndoc["footnotes"]),
        "toc_links": n_links, "missing_images": ndoc["missing_images"],
    }


def export_document(doc: Any, out_dir: str | Path, stem: str = "report", *,
                    formats: tuple[str, ...] = ("html", "pdf", "docx"),
                    html_mode: str = "self", assets_root: Path | None = None,
                    log: Callable[[str], None] = print) -> dict:
    """세 형식을 한 번에 굽는다 — {형식: 산출 요약}."""
    out_dir = Path(out_dir)
    result: dict[str, dict] = {}
    if "html" in formats:
        result["html"] = export_html(doc, out_dir / f"{stem}.html", mode=html_mode,
                                     assets_root=assets_root, log=log)
    if "pdf" in formats:
        result["pdf"] = export_pdf(doc, out_dir / f"{stem}.pdf",
                                   assets_root=assets_root, log=log)
    if "docx" in formats:
        result["docx"] = export_docx(doc, out_dir / f"{stem}.docx",
                                     assets_root=assets_root, log=log)
    return result
